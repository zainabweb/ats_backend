"""
Raw text extraction from PDF / DOCX. No chunking — the full document text
is returned as a single string, which is what gets embedded as ONE vector.
"""
import os

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from langdetect import DetectorFactory, LangDetectException, detect

from app.core.exceptions import ResumeParseError
from app.core.logging import get_logger

logger = get_logger(__name__)

# Deterministic detection - langdetect is otherwise seeded from wall-clock
# time, which would make the same PDF classify differently between runs.
DetectorFactory.seed = 0


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(file_path)
        if ext == ".docx":
            return _extract_docx(file_path)
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError(f"Could not read {os.path.basename(file_path)}: {exc}") from exc

    raise ResumeParseError(f"Unsupported file extension: {ext}")


def _extract_pdf(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text_parts.append(_extract_page_text(page))
    text_parts = _select_primary_language_pages(text_parts)
    text = "\n".join(text_parts).strip()
    if not text:
        raise ResumeParseError("PDF contains no extractable text (possibly a scanned image).")
    return text


def _detect_page_lang(text: str) -> str | None:
    """Best-effort language code for one page. Returns None when the page
    is too short/sparse (a lone header, a contact strip, a blank divider)
    for langdetect to be reliable - those pages get attached to whichever
    version they physically sit inside, in _select_primary_language_pages,
    rather than forcing a language of their own."""
    stripped = text.strip()
    if len(stripped) < 30:
        return None
    try:
        return detect(stripped)
    except LangDetectException:
        return None


def _select_primary_language_pages(page_texts: list[str]) -> list[str]:
    """
    Some CV PDFs bundle more than one full copy of the same resume back to
    back - most commonly an English copy followed by a translated (e.g.
    Urdu) copy meant for a different reader. Concatenating every page
    blindly (the old behaviour) hands the LLM two overlapping-but-not-
    identical accounts of the same person in one blob, which is exactly
    what causes skills/experience/education to be read from whichever copy
    happened to phrase something more clearly, or duplicated across both.

    Detect each page's language, fill short/undetectable pages in with
    whichever language surrounds them (so a one-line section header doesn't
    get treated as its own "version"), and - only when more than one real
    language is present AND English is one of them - keep just the English
    pages, matching what the user actually wants scored. A resume that's a
    single language throughout (the overwhelming common case), or one with
    no English copy at all, is returned completely untouched.
    """
    if len(page_texts) <= 1:
        return page_texts

    raw_langs = [_detect_page_lang(t) for t in page_texts]

    # Forward-fill, then back-fill, undetectable pages from their nearest
    # detected neighbour, so a short page is treated as part of whichever
    # version physically surrounds it.
    filled: list[str | None] = list(raw_langs)
    last_known = None
    for i, lang in enumerate(filled):
        if lang is not None:
            last_known = lang
        elif last_known is not None:
            filled[i] = last_known
    next_known = None
    for i in range(len(filled) - 1, -1, -1):
        if filled[i] is not None:
            next_known = filled[i]
        elif next_known is not None:
            filled[i] = next_known

    distinct = {l for l in filled if l is not None}
    if len(distinct) <= 1:
        return page_texts  # single language throughout (or nothing detectable at all)

    if "en" not in distinct:
        return page_texts  # multiple languages, but no English copy to prefer

    logger.info(
        "multiple_language_versions_detected_preferring_english",
        extra={"page_languages": filled},
    )
    return [t for t, lang in zip(page_texts, filled) if lang == "en"]


def _extract_page_text(page) -> str:
    """
    Multi-column resumes (a narrow sidebar - Address/Phone/Email - next to a
    wider main column - Experience/Education) break plain extract_text(): it
    lays text out strictly by vertical position across the FULL page width,
    so a sidebar line and a main-column line at the same height get merged/
    interleaved into one garbled sequence - which is exactly what separates
    a date range from the role title that should be directly below it.

    Fix: detect a wide vertical gap with no words in it (a column gutter),
    split words left/right of that gap, and extract each column fully,
    top-to-bottom, on its own - so a date-then-title block inside one column
    is never interrupted by the other column's lines. Falls back to the
    normal single-pass extraction when no such gap exists, so single-column
    resumes are completely unaffected.
    """
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    if not words:
        return page.extract_text() or ""

    page_width = page.width
    slice_width = 5
    n_slices = max(1, int(page_width // slice_width) + 1)
    covered = [False] * n_slices
    for w in words:
        start = int(w["x0"] // slice_width)
        end = int(w["x1"] // slice_width)
        for i in range(max(0, start), min(n_slices, end + 1)):
            covered[i] = True

    min_gap_slices = max(1, int(20 // slice_width))  # ~20pt minimum gutter
    gaps = []
    run_start = None
    for i, is_covered in enumerate(covered):
        if not is_covered and run_start is None:
            run_start = i
        elif is_covered and run_start is not None:
            if i - run_start >= min_gap_slices:
                gaps.append((run_start, i))
            run_start = None
    if run_start is not None and n_slices - run_start >= min_gap_slices:
        gaps.append((run_start, n_slices))

    # Only treat it as a real column boundary if it sits away from the page
    # edges (not just trailing whitespace at the end of a line).
    column_bounds = [
        gap for gap in gaps
        if 0.15 * page_width < gap[0] * slice_width < 0.85 * page_width
    ]

    if not column_bounds:
        return page.extract_text() or ""

    g = column_bounds[0]
    split_x = g[0] * slice_width + (g[1] - g[0]) * slice_width / 2

    left_words = [w for w in words if w["x0"] < split_x]
    right_words = [w for w in words if w["x0"] >= split_x]

    def _words_to_lines(ws):
        lines = {}
        for w in ws:
            key = round(w["top"] / 3)  # group words on (roughly) the same line
            lines.setdefault(key, []).append(w)
        out = []
        for key in sorted(lines):
            row = sorted(lines[key], key=lambda w: w["x0"])
            out.append((key, " ".join(w["text"] for w in row)))
        return out

    # Many resumes use the narrow column purely as a strip of SECTION LABELS
    # (Experience / Volunteer / Education, or Objective / References /
    # Address) that sit beside the actual entries in the wide column, rather
    # than two independent streams of content. Dumping the whole left column
    # first and the whole right column after (the old behaviour) rips every
    # label away from the content it heads - the LLM then sees one
    # undifferentiated blob of roles/education/volunteer work with no
    # section boundaries at all, which is exactly what causes education or
    # volunteer entries to get counted as paid experience, or several roles
    # to be read as one. Instead, merge both columns back into a single
    # top-to-bottom stream ordered by vertical position (left first on a
    # shared line), so each label lands right next to the content at that
    # height - the closest we can get to the resume's real reading order.
    left_lines = _words_to_lines(left_words)
    right_lines = _words_to_lines(right_words)
    merged = sorted(
        [(key, 0, text) for key, text in left_lines]
        + [(key, 1, text) for key, text in right_lines]
    )
    return "\n".join(text for _, _, text in merged)


def _paragraph_full_text(p_element) -> str:
    """
    Text of one <w:p>, INCLUDING runs nested inside a content control
    (<w:sdt>/<w:sdtContent>) - a field Word inserts for "click here to add
    a date" style placeholders. python-docx's Paragraph.text only looks at
    <w:r> elements that are direct children of <w:p>; a run one level
    deeper inside a content control is invisible to it. Resume templates
    commonly use exactly this field for each role's date range, so relying
    on Paragraph.text silently drops the date while keeping the job title
    right next to it - which is what makes a role with a real, filled-in
    date look date-less. Walking every descendant <w:t> instead (regardless
    of nesting depth) picks up both plain runs and content-control runs.
    """
    return "".join(t.text or "" for t in p_element.iter(qn("w:t")))


def _cell_full_text(cell) -> str:
    return "\n".join(_paragraph_full_text(p._p) for p in cell.paragraphs)


def _extract_docx(file_path: str) -> str:
    """
    Many resumes are built as ONE table used purely for visual layout
    (a sidebar column, section-header bands, etc.), with cells merged both
    across columns (gridSpan) and down rows (vMerge) for that look. Word/
    python-docx exposes a merged cell at EVERY grid position it visually
    covers, returning the same text each time - and some templates carry
    that further by literally repeating the same paragraph content in every
    row a vMerge spans, rather than leaving continuation rows empty. Walking
    every row/cell naively (the old behaviour) re-appends that one physical
    block of text once per position it covers, so a name/contact block or a
    role can end up duplicated 4-12x in what the LLM reads - which crowds
    out the real structure of the resume and makes reliable section
    extraction much harder. Skip a cell already captured for this table,
    either by its underlying XML element (handles gridSpan/vMerge sharing
    the same node) or by exact text (handles templates that duplicate the
    text itself into every spanned row) - the same content is then read
    into the pipeline once, at its first occurrence, matching the resume's
    actual reading order.
    """
    doc = Document(file_path)
    paragraphs = [
        text for p in doc.paragraphs if (text := _paragraph_full_text(p._p).strip())
    ]
    for table in doc.tables:
        # Keep the actual XML elements (not id()) in the seen-set: python-docx
        # builds a fresh _Cell wrapper on every `row.cells` access, so a
        # wrapper (and the id() of its _tc) can be garbage-collected and that
        # same memory address handed to a LATER, genuinely different cell -
        # id()-based dedup would then wrongly treat that unrelated cell as
        # already seen and drop its content. Holding the element itself keeps
        # it alive and compares by real identity for the whole table.
        seen_tc_elements: set = set()
        seen_text: set[str] = set()
        for row in table.rows:
            for cell in row.cells:
                text = _cell_full_text(cell).strip()
                if not text:
                    continue
                if cell._tc in seen_tc_elements or text in seen_text:
                    continue
                seen_tc_elements.add(cell._tc)
                seen_text.add(text)
                paragraphs.append(text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ResumeParseError("DOCX contains no extractable text.")
    return text
